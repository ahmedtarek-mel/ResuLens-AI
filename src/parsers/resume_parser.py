"""
Resume Parser Module
Extracts structured information from PDF and DOCX resume files
"""

import re
import io
from pathlib import Path
from typing import Dict, List, Any, Optional, Union
from dataclasses import dataclass, field

import pdfplumber
from docx import Document


@dataclass
class ParsedResume:
    """Structured resume data"""
    raw_text: str
    contact_info: Dict[str, str] = field(default_factory=dict)
    summary: str = ""
    experience: List[Dict[str, Any]] = field(default_factory=list)
    education: List[Dict[str, Any]] = field(default_factory=list)
    skills: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)
    projects: List[Dict[str, Any]] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "raw_text": self.raw_text,
            "contact_info": self.contact_info,
            "summary": self.summary,
            "experience": self.experience,
            "education": self.education,
            "skills": self.skills,
            "certifications": self.certifications,
            "projects": self.projects
        }


class ResumeParser:
    """
    Multi-format resume parser with section extraction
    Supports PDF, DOCX, and plain text files
    """
    
    # Section headers patterns
    SECTION_PATTERNS = {
        'summary': r'(?i)(summary|objective|profile|about\s*me)',
        'experience': r'(?i)(experience|work\s*history|employment|professional\s*experience)',
        'education': r'(?i)(education|academic|qualifications|degrees)',
        'skills': r'(?i)(skills|technical\s*skills|competencies|technologies|expertise)',
        'certifications': r'(?i)(certifications?|certificates?|licenses?|credentials)',
        'projects': r'(?i)(projects?|portfolio|personal\s*projects)'
    }
    
    # Contact extraction patterns
    EMAIL_PATTERN = r'[\w\.-]+@[\w\.-]+\.\w+'
    PHONE_PATTERN = r'[\+]?[(]?[0-9]{3}[)]?[-\s\.]?[(]?[0-9]{3}[)]?[-\s\.]?[0-9]{4,6}'
    LINKEDIN_PATTERN = r'(?:linkedin\.com/in/|linkedin:?\s*)([a-zA-Z0-9-]+)'
    GITHUB_PATTERN = r'(?:github\.com/|github:?\s*)([a-zA-Z0-9-]+)'
    
    def parse(self, file: Union[str, Path, io.BytesIO], filename: str = None) -> ParsedResume:
        """
        Parse a resume file and extract structured information
        
        Args:
            file: File path or BytesIO object
            filename: Original filename (needed for BytesIO to determine type)
            
        Returns:
            ParsedResume object with extracted data
        """
        # Determine file type and extract text
        if isinstance(file, (str, Path)):
            file_path = Path(file)
            extension = file_path.suffix.lower()
            raw_text = self._extract_text_from_path(file_path, extension)
        else:
            # BytesIO object
            extension = Path(filename).suffix.lower() if filename else '.pdf'
            raw_text = self._extract_text_from_bytes(file, extension)
        
        # Parse the extracted text
        return self._parse_text(raw_text)
    
    def _extract_text_from_path(self, file_path: Path, extension: str) -> str:
        """Extract text from a file path"""
        if extension == '.pdf':
            return self._extract_from_pdf_path(file_path)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx_path(file_path)
        elif extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _extract_text_from_bytes(self, file_bytes: io.BytesIO, extension: str) -> str:
        """Extract text from bytes"""
        file_bytes.seek(0)
        
        if extension == '.pdf':
            return self._extract_from_pdf_bytes(file_bytes)
        elif extension in ['.docx', '.doc']:
            return self._extract_from_docx_bytes(file_bytes)
        elif extension == '.txt':
            return file_bytes.read().decode('utf-8')
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _extract_from_pdf_path(self, file_path: Path) -> str:
        """Extract text from PDF file"""
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts)
    
    def _extract_from_pdf_bytes(self, file_bytes: io.BytesIO) -> str:
        """Extract text from PDF bytes"""
        text_parts = []
        with pdfplumber.open(file_bytes) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n'.join(text_parts)
    
    def _extract_from_docx_path(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    def _extract_from_docx_bytes(self, file_bytes: io.BytesIO) -> str:
        """Extract text from DOCX bytes"""
        doc = Document(file_bytes)
        return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    def _parse_text(self, raw_text: str) -> ParsedResume:
        """Parse raw text into structured sections"""
        resume = ParsedResume(raw_text=raw_text)
        
        # Extract contact information
        resume.contact_info = self._extract_contact_info(raw_text)
        
        # Split into sections
        sections = self._split_into_sections(raw_text)
        
        # Extract each section
        resume.summary = self._clean_text(sections.get('summary', ''))
        resume.experience = self._parse_experience(sections.get('experience', ''))
        resume.education = self._parse_education(sections.get('education', ''))
        resume.skills = self._parse_skills(sections.get('skills', ''), raw_text)
        resume.certifications = self._parse_certifications(sections.get('certifications', ''))
        resume.projects = self._parse_projects(sections.get('projects', ''))
        
        return resume
    
    def _extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information from text"""
        contact = {}
        
        # Email
        email_match = re.search(self.EMAIL_PATTERN, text)
        if email_match:
            contact['email'] = email_match.group()
        
        # Phone
        phone_match = re.search(self.PHONE_PATTERN, text)
        if phone_match:
            contact['phone'] = phone_match.group()
        
        # LinkedIn
        linkedin_match = re.search(self.LINKEDIN_PATTERN, text)
        if linkedin_match:
            contact['linkedin'] = linkedin_match.group(1)
        
        # GitHub
        github_match = re.search(self.GITHUB_PATTERN, text)
        if github_match:
            contact['github'] = github_match.group(1)
        
        # Name (usually first line or most prominent text)
        lines = text.split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            # Heuristic: name is usually 2-4 words, all caps or title case
            if line and len(line.split()) <= 4 and not re.search(r'[@\d]', line):
                if not any(keyword in line.lower() for keyword in ['resume', 'cv', 'curriculum']):
                    contact['name'] = line
                    break
        
        return contact
    
    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """Split resume text into sections based on headers"""
        sections = {}
        lines = text.split('\n')
        
        current_section = 'header'
        current_content = []
        
        for line in lines:
            line_stripped = line.strip()
            
            # Check if this line is a section header
            found_section = None
            for section, pattern in self.SECTION_PATTERNS.items():
                if re.match(pattern, line_stripped):
                    found_section = section
                    break
            
            if found_section:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                
                # Start new section
                current_section = found_section
                current_content = []
            else:
                current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    
    def _parse_experience(self, text: str) -> List[Dict[str, Any]]:
        """Parse work experience section"""
        experiences = []
        if not text.strip():
            return experiences
        
        # Split by common patterns (dates, company names followed by titles)
        date_pattern = r'(\d{4}\s*[-–]\s*(?:\d{4}|present|current)|\d{1,2}/\d{4}\s*[-–]\s*(?:\d{1,2}/\d{4}|present|current))'
        
        chunks = re.split(date_pattern, text, flags=re.IGNORECASE)
        
        current_exp = {}
        for i, chunk in enumerate(chunks):
            chunk = chunk.strip()
            if not chunk:
                continue
            
            # Check if this is a date
            if re.match(date_pattern, chunk, re.IGNORECASE):
                if current_exp:
                    experiences.append(current_exp)
                current_exp = {'dates': chunk, 'description': ''}
            elif current_exp:
                # This is description for current experience
                lines = chunk.split('\n')
                if lines and not current_exp.get('title'):
                    current_exp['title'] = lines[0].strip()
                    current_exp['description'] = '\n'.join(lines[1:]).strip()
                else:
                    current_exp['description'] += '\n' + chunk
        
        if current_exp:
            experiences.append(current_exp)
        
        return experiences
    
    def _parse_education(self, text: str) -> List[Dict[str, Any]]:
        """Parse education section"""
        education = []
        if not text.strip():
            return education
        
        # Common degree patterns
        degree_pattern = r'(?i)(bachelor|master|phd|ph\.d|b\.s\.|m\.s\.|b\.a\.|m\.a\.|bsc|msc|mba|associate)'
        
        lines = text.split('\n')
        current_edu = {}
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if re.search(degree_pattern, line):
                if current_edu:
                    education.append(current_edu)
                current_edu = {'degree': line, 'details': ''}
            elif current_edu:
                current_edu['details'] += line + ' '
        
        if current_edu:
            education.append(current_edu)
        
        return education
    
    def _parse_skills(self, skills_section: str, full_text: str) -> List[str]:
        """Extract skills from skills section and full text"""
        skills = set()
        
        # Process skills section
        if skills_section:
            # Split by common delimiters
            skill_parts = re.split(r'[,;•|\n]', skills_section)
            for part in skill_parts:
                skill = self._clean_skill(part)
                if skill and len(skill) > 1 and len(skill) < 50:
                    skills.add(skill)
        
        return list(skills)
    
    def _clean_skill(self, skill: str) -> str:
        """Clean and normalize a skill string"""
        skill = skill.strip()
        # Remove bullets and special characters at start
        skill = re.sub(r'^[\s•\-\*\d\.]+', '', skill)
        # Remove trailing punctuation
        skill = re.sub(r'[:\s]+$', '', skill)
        return skill.strip()
    
    def _parse_certifications(self, text: str) -> List[str]:
        """Parse certifications section"""
        certifications = []
        if not text.strip():
            return certifications
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if line and len(line) > 3:
                # Clean up the certification name
                cert = re.sub(r'^[\s•\-\*\d\.]+', '', line)
                if cert:
                    certifications.append(cert)
        
        return certifications
    
    def _parse_projects(self, text: str) -> List[Dict[str, Any]]:
        """Parse projects section"""
        projects = []
        if not text.strip():
            return projects
        
        # Split by bullet points or numbered lists
        project_chunks = re.split(r'\n(?=[\s]*[•\-\*\d\.])', text)
        
        for chunk in project_chunks:
            chunk = chunk.strip()
            if chunk and len(chunk) > 10:
                lines = chunk.split('\n')
                project = {
                    'title': re.sub(r'^[\s•\-\*\d\.]+', '', lines[0]).strip(),
                    'description': '\n'.join(lines[1:]).strip() if len(lines) > 1 else ''
                }
                if project['title']:
                    projects.append(project)
        
        return projects
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove leading/trailing whitespace
        text = text.strip()
        return text
